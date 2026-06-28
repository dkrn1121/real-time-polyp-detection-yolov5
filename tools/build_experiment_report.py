from __future__ import annotations

import argparse
import json
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_GITHUB_URL = "https://github.com/dkrn1121/real-time-polyp-detection-yolov5"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    set_run_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    set_run_font(p.add_run(text))
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, color="2E74B5", bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9, color="555555")
    return p


def add_key_value_table(doc, rows, widths=(1.6, 4.9)):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths)
    table.rows[0].cells[0].text = rows[0][0]
    table.rows[0].cells[1].text = rows[0][1]
    for row_data in rows[1:]:
        cells = table.add_row().cells
        cells[0].text = row_data[0]
        cells[1].text = row_data[1]
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=10.5)
        set_cell_shading(row.cells[0], "F2F4F7")
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    return table


def add_data_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=10.5, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
            cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in cells[idx].paragraphs[0].runs:
                set_run_font(run, size=10.5)
    return table


def add_image(doc, rel_path, caption, width=6.2):
    image_path = ROOT / rel_path
    if not image_path.exists():
        add_paragraph(doc, f"图像文件缺失: {rel_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = "YOLOv5 息肉实时检测实验报告"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color="555555")

    footer = section.footer.paragraphs[0]
    footer.text = "基于公开医学数据集的目标检测课程作业"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color="555555")


def build_report(output_path: Path, github_url: str, github_status: str):
    metrics = json.loads((ROOT / "evaluation_results" / "metrics.json").read_text(encoding="utf-8"))
    fps = json.loads((ROOT / "evaluation_results" / "fps_benchmark.json").read_text(encoding="utf-8"))

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(48)
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run("基于 YOLOv5 的实时息肉目标检测实验报告")
    set_run_font(title_run, size=24, color="0B2545", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    set_run_font(subtitle.add_run("Real-time Polyp Detection with YOLOv5n"), size=13, color="555555")

    add_key_value_table(
        doc,
        [
            ("课程作业", "基于 YOLOv5 的实时目标检测"),
            ("姓名", "__________"),
            ("学号", "__________"),
            ("班级", "__________"),
            ("实验日期", "2026 年 6 月 28 日"),
            ("项目目录", str(ROOT)),
        ],
    )
    doc.add_page_break()

    add_heading(doc, "一、实验背景与目的", 1)
    add_paragraph(
        doc,
        "目标检测任务需要同时给出目标类别和边界框位置。医学内镜场景中，息肉目标通常面积较小，"
        "且容易受到反光、遮挡、运动模糊和黏膜纹理相似性的影响，因此对检测模型的定位能力和实时性都有较高要求。"
    )
    add_paragraph(
        doc,
        "本实验基于公开息肉检测数据集，训练轻量化 YOLOv5n 模型，实现对内镜图像和视频帧中的息肉目标进行实时检测。"
        "实验重点包括模型训练、mAP 指标评估、FPS 推理速度测试、检测结果可视化以及失败案例原因分析。"
    )
    add_paragraph(
        doc,
        "从应用背景看，结直肠息肉是临床内镜筛查中需要重点关注的异常区域之一。传统人工阅片依赖医生经验，"
        "在长时间检查、病灶较小或画面质量不稳定的情况下，可能受到注意力下降和主观判断差异的影响。"
        "基于深度学习的目标检测方法能够在视频流中持续给出候选区域，为辅助筛查、教学演示和离线质控提供技术基础。"
    )
    add_paragraph(
        doc,
        "与通用机器视觉场景相比，医学内镜检测的难点更加集中。第一，息肉通常与周围黏膜颜色和纹理相近，"
        "类别判别边界不如行人、车辆等自然图像目标明显；第二，内镜图像存在圆形视野、暗角、镜头反光、液体遮挡和器械干扰；"
        "第三，视频帧之间存在快速运动和局部模糊，要求模型既有较高的召回率，又不能牺牲实时性。"
        "因此，本实验选择轻量化 YOLOv5n 作为基础模型，在精度和速度之间取得平衡。"
    )

    add_heading(doc, "二、数据集介绍", 1)
    add_key_value_table(
        doc,
        [
            ("数据来源", "Roboflow Universe: bsc-train_psc1toc5 Version 7"),
            ("数据许可", "CC BY 4.0"),
            ("检测类别", "polyp"),
            ("训练集", "2166 张图像及标签"),
            ("验证集", "545 张图像及标签"),
            ("本地配置", "polyp_dataset/data.yaml"),
        ],
    )
    add_paragraph(
        doc,
        "数据标注格式为 YOLO 检测格式，标签中包含类别编号、目标中心点坐标、宽度和高度。"
        "本实验只检测一个医学目标类别，因此类别数 nc=1，类别名称为 polyp。"
    )

    add_heading(doc, "三、理论基础与模型方法", 1)
    add_paragraph(
        doc,
        "目标检测任务本质上由分类与定位两部分组成：分类分支判断候选区域属于哪一类目标，定位分支回归边界框坐标。"
        "YOLO 系列方法将目标检测建模为单阶段回归问题，在一次前向传播中同时输出目标类别、置信度和边界框位置，"
        "相较两阶段检测器具有更低的推理延迟，适合视频实时检测场景。"
    )
    add_paragraph(
        doc,
        "YOLOv5 的整体结构可概括为 Backbone、Neck 和 Head 三个部分。Backbone 负责从输入图像中提取多尺度语义特征；"
        "Neck 通过特征融合增强浅层纹理信息与深层语义信息的表达；Head 在不同尺度上预测边界框、目标置信度和类别概率。"
        "这种多尺度检测结构对息肉这类大小变化明显、边缘形态不规则的医学目标较为重要。"
    )
    add_paragraph(
        doc,
        "模型训练过程中，损失函数通常由边界框回归损失、目标置信度损失和类别损失组成。边界框回归用于衡量预测框与真实框的重合程度，"
        "类别损失用于约束类别判断，置信度损失用于区分前景目标与背景区域。对于单类别息肉检测而言，类别判断相对简单，"
        "但边界框定位质量会直接影响 mAP@0.5:0.95 等严格指标，因此边界细节仍是模型性能提升的关键。"
    )
    add_paragraph(
        doc,
        "本实验采用 YOLOv5n 版本，是 YOLOv5 系列中的轻量模型。其参数量和计算量较小，适合在普通 CPU 或低算力设备上验证实时性。"
        "虽然轻量模型的表达能力弱于更大的 YOLOv5s、YOLOv5m 等版本，但在课堂实验和实时演示场景中更能体现速度优势。"
    )

    add_heading(doc, "四、实验环境与模型配置", 1)
    add_data_table(
        doc,
        ["项目", "设置"],
        [
            ("模型", "YOLOv5n / yolov5nu.pt"),
            ("训练权重", "runs/detect/train/weights/best.pt"),
            ("输入尺寸", "640 x 640"),
            ("训练轮数", "100 epochs"),
            ("Batch size", "8"),
            ("优化器", "Ultralytics auto"),
            ("推理设备", fps["device"]),
            ("视频规格", "480 x 360, 30 FPS, 253 帧"),
        ],
        (2.1, 4.4),
    )
    add_paragraph(
        doc,
        "训练中启用基础增强，包括上下翻转、左右翻转、HSV 饱和度和亮度扰动；Mosaic、Mixup 和 Copy-Paste 增强保持关闭。"
        "该设置更适合医学图像中较稳定的组织结构，避免过强拼接增强破坏真实内镜纹理。"
    )

    add_heading(doc, "五、实验设计与技术路线", 1)
    add_paragraph(
        doc,
        "实验整体采用“数据准备、模型训练、定量评估、视频推理、结果分析”的技术路线。"
        "数据准备阶段重点保证图像、标签与类别配置一致；模型训练阶段利用预训练权重进行迁移学习，"
        "使模型在有限医学数据上更快获得稳定特征表示；定量评估阶段从准确性和实时性两个维度衡量模型性能；"
        "可视化分析阶段结合验证集预测图、混淆矩阵、PR 曲线和视频帧截图，对模型优势与不足进行解释。"
    )
    add_paragraph(
        doc,
        "在训练策略上，本实验没有盲目追求复杂增强，而是采用较克制的数据增强设置。医学图像具有较强的结构真实性，"
        "过强的拼接或混合增强可能破坏内镜纹理和病灶边界，因此实验保留翻转、亮度和饱和度扰动等较自然的变化，"
        "同时关闭 Mosaic、Mixup 和 Copy-Paste 等可能引入不真实组合的增强方式。该设计有助于模型学习更接近真实内镜场景的特征。"
    )
    add_paragraph(
        doc,
        "在评估策略上，单一指标不足以完整描述模型性能。Precision 反映误检控制能力，Recall 反映漏检控制能力，"
        "mAP@0.5 关注较宽松 IoU 阈值下的检测成功率，mAP@0.5:0.95 则进一步考察边界框精细定位能力。"
        "对于医学辅助检测而言，漏检通常比少量误检更值得关注，因此在分析时需要综合 Precision、Recall 与 mAP 曲线，"
        "而不是只依据某一个最高数值判断模型优劣。"
    )
    add_paragraph(
        doc,
        "在实时性评估上，实验使用同一段内镜视频逐帧推理，分别统计普通检测和带目标跟踪的平均耗时。"
        "这样可以区分纯检测模型的速度上限与实际视频分析场景中的处理开销，使 FPS 结果更接近真实应用。"
    )

    add_image(doc, "runs/detect/train/results.png", "图 1 训练损失、Precision、Recall 和 mAP 曲线", 6.3)

    add_heading(doc, "六、实验结果与指标分析", 1)
    add_heading(doc, "6.1 验证集检测精度", 2)
    add_data_table(
        doc,
        ["指标", "数值", "说明"],
        [
            ("Precision", f"{metrics['precision']:.4f}", "预测为息肉的目标中，正确预测比例较高"),
            ("Recall", f"{metrics['recall']:.4f}", "真实息肉中被模型检出的比例较高"),
            ("F1-Score", f"{metrics['f1_score']:.4f}", "Precision 与 Recall 的综合表现"),
            ("mAP@0.5", f"{metrics['map50']:.4f}", "IoU=0.5 时定位性能较好"),
            ("mAP@0.5:0.95", f"{metrics['map50_95']:.4f}", "严格 IoU 阈值下仍有边界优化空间"),
        ],
        (1.55, 1.1, 3.85),
    )
    add_paragraph(
        doc,
        "从指标上看，模型 Precision 为 0.8581，说明被模型判定为息肉的目标中，大部分与真实标注一致，误检率处于可接受范围；"
        "Recall 为 0.8322，说明模型能够发现验证集中多数息肉目标，但仍存在一定漏检。F1-Score 为 0.8450，"
        "表明 Precision 与 Recall 之间保持了较好的平衡，没有出现只追求高置信度而牺牲召回率，或过度追求召回而导致大量误检的情况。"
    )
    add_paragraph(
        doc,
        "mAP@0.5 达到 0.8913，说明在 IoU=0.5 的常用阈值下，模型对息肉目标具有较好的整体检测能力。"
        "但 mAP@0.5:0.95 为 0.6791，低于 mAP@0.5，说明当评价标准逐步提高、要求预测框与真实框更精确重合时，"
        "模型仍会受到边界不清晰、目标形态变化和图像噪声的影响。这一差异符合医学小目标检测的典型特点：模型通常能够发现目标大致位置，"
        "但要稳定贴合病灶边缘仍更困难。"
    )
    add_image(doc, "runs/detect/val-4/confusion_matrix.png", "图 2 验证集混淆矩阵", 5.4)
    add_image(doc, "runs/detect/val-4/BoxPR_curve.png", "图 3 Precision-Recall 曲线", 5.6)
    add_image(doc, "runs/detect/val-4/BoxF1_curve.png", "图 4 F1-Confidence 曲线", 5.6)

    add_heading(doc, "6.2 推理速度", 2)
    add_data_table(
        doc,
        ["模式", "平均耗时", "FPS", "备注"],
        [
            (
                "detect",
                f"{fps['detect']['avg_ms_per_frame']:.2f} ms/帧",
                f"{fps['detect']['fps']:.2f}",
                "普通检测推理",
            ),
            (
                "track",
                f"{fps['track']['avg_ms_per_frame']:.2f} ms/帧",
                f"{fps['track']['fps']:.2f}",
                "带目标 ID 的视频跟踪",
            ),
        ],
        (1.3, 1.7, 1.1, 2.4),
    )
    add_paragraph(
        doc,
        f"FPS 测试在 {fps['device']} 上完成，输入视频共 {int(fps['video']['frame_count'])} 帧，"
        f"脚本预热 {fps['warmup_frames']} 帧后计时。detect 模式达到 {fps['detect']['fps']:.2f} FPS，"
        f"track 模式达到 {fps['track']['fps']:.2f} FPS，基本满足课堂作业中实时检测的要求。"
    )
    add_paragraph(
        doc,
        "从实时性角度看，普通检测模式的速度略高于跟踪模式，原因是跟踪需要额外维护目标 ID、关联连续帧中的检测框，"
        "会带来一定计算开销。即便如此，track 模式仍接近视频原始帧率，说明 YOLOv5n 的轻量结构适合实时视频分析。"
        "在更高分辨率或更复杂模型下，FPS 可能下降，因此实时检测系统通常需要在图像尺寸、模型规模和硬件平台之间做折中。"
    )

    add_heading(doc, "七、检测结果可视化", 1)
    add_image(doc, "runs/detect/val-4/val_batch0_pred.jpg", "图 5 验证集批量预测示例", 6.3)
    add_image(doc, "report/0.23.jpg", "图 6 视频第 0.23 秒检测结果", 4.8)
    add_image(doc, "report/6.63.jpg", "图 7 视频第 6.63 秒检测与跟踪结果", 4.8)
    add_image(doc, "report/7.16.jpg", "图 8 视频第 7.16 秒检测与跟踪结果", 4.8)
    add_paragraph(
        doc,
        "从视频帧结果可以看到，模型能够在内镜画面中给出息肉边界框、类别名称和置信度；track 模式还会为连续帧中的同一目标分配 ID，"
        "便于观察目标在视频中的位置变化。"
    )
    add_paragraph(
        doc,
        "验证集批量预测图显示，模型在目标较明显、边界较清晰的样本上能够给出较高置信度的检测框；"
        "视频截图则体现了模型在连续帧中的应用效果。相比静态图像，视频帧还包含运动模糊、视角变化和局部遮挡，"
        "因此可视化结果不仅用于展示模型效果，也用于判断模型在真实动态场景中的稳定性。"
    )

    add_heading(doc, "八、失败案例与原因分析", 1)
    add_bullet(doc, "高光和反光区域会造成局部纹理异常，模型可能把强反光或器械边缘误判为息肉。")
    add_bullet(doc, "息肉与正常黏膜颜色接近时，边界框可能偏大或偏小，导致 mAP@0.5:0.95 低于 mAP@0.5。")
    add_bullet(doc, "目标出现在画面边缘、被器械遮挡或面积较小时，模型容易出现漏检。")
    add_bullet(doc, "数据集为单类别检测，负样本和复杂病例覆盖仍有限，泛化到更多医院或设备的视频时需要继续验证。")
    add_paragraph(
        doc,
        "改进方向包括扩充不同设备和不同光照条件下的数据、增加难例负样本、尝试更高容量模型或目标检测与分割联合训练，"
        "并在真实视频流中进一步测试延迟和稳定性。"
    )
    add_paragraph(
        doc,
        "从误差来源看，漏检通常发生在目标面积较小、边缘不完整或颜色与背景高度接近的区域；误检则多与高光、黏液、器械和褶皱有关。"
        "这说明模型学习到的纹理特征仍可能与非病灶结构发生混淆。若要进一步提高临床场景下的鲁棒性，需要引入更多跨设备、跨患者、"
        "跨光照条件的数据，并在标注阶段尽可能统一边界定义，减少训练标签本身的不确定性。"
    )

    add_heading(doc, "九、GitHub 获取项目代码", 1)
    if github_status == "pushed":
        status_text = "项目代码已整理并推送到 GitHub，可通过以下方式获取："
    else:
        status_text = "项目默认整理为以下 GitHub 地址；若远端仓库已创建并完成推送，可通过以下方式获取："
    add_paragraph(doc, status_text)
    add_paragraph(doc, github_url)
    add_paragraph(doc, f"git clone {github_url}.git")
    add_paragraph(doc, f"cd {Path(github_url).name}")
    add_paragraph(doc, "pip install -r requirements.txt")
    add_paragraph(
        doc,
        "为避免泄露密钥和造成仓库过大，完整 polyp_dataset 数据集、.env 文件和冗余训练缓存不纳入普通 Git 提交；"
        "Roboflow API Key 通过 ROBOFLOW_API_KEY 环境变量读取。"
    )

    add_heading(doc, "十、实验结论", 1)
    add_paragraph(
        doc,
        "本实验完成了基于 YOLOv5n 的息肉单类别目标检测训练与验证。模型在验证集上取得 mAP@0.5=0.8913，"
        "CPU 条件下 detect 模式达到 29 FPS 左右，说明轻量化 YOLOv5n 在该数据集上具备较好的实时检测能力。"
    )
    add_paragraph(
        doc,
        "同时，严格 IoU 下的 mAP@0.5:0.95 为 0.6791，说明模型对小目标边界细节仍有提升空间。"
        "后续可从数据扩充、难例增强、模型结构和医学场景验证等方向进一步改进。"
    )
    add_paragraph(
        doc,
        "总体而言，本实验验证了单阶段目标检测模型在医学内镜息肉检测任务中的可行性。YOLOv5n 能够在较低计算开销下完成目标定位，"
        "适合作为实时检测系统的基础模型；但医学图像任务对可靠性要求较高，后续如果面向更严肃的应用场景，还需要进一步开展多中心数据验证、"
        "失败案例复盘和模型可解释性分析。"
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def parse_args():
    desktop = Path(os.environ.get("USERPROFILE", str(ROOT))) / "Desktop"
    parser = argparse.ArgumentParser(description="Build YOLOv5 experiment report DOCX.")
    parser.add_argument("--output", default=str(desktop / "YOLOv5息肉实时检测实验报告.docx"))
    parser.add_argument("--github-url", default=DEFAULT_GITHUB_URL)
    parser.add_argument("--github-status", choices=["pending", "pushed"], default="pending")
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build_report(Path(args.output), args.github_url, args.github_status)
    print(args.output)
